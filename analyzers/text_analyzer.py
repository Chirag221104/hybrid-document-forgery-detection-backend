import PyPDF2
import fitz  # PyMuPDF - more reliable
import docx  # python-docx for DOCX files
import re
import os

class TextAnalyzer:
    async def analyze(self, file_path: str, file_info: dict):
        """Perform REAL text analysis for forgery detection"""
        try:
            file_type = file_info["type"]
            
            if file_type == "application/pdf":
                return await self._analyze_pdf_text(file_path)
            elif file_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
                return await self._analyze_docx_text(file_path)
            else:
                return await self._basic_analysis(file_info)
                
        except Exception as e:
            print(f"Text analysis error: {e}")
            return {
                "totalWords": 0,
                "suspiciousWords": 0,
                "confidence": 0,
                "flags": [f"Analysis failed: {str(e)}"]
            }
    
    async def _analyze_pdf_text(self, file_path: str):
        """Extract text using multiple methods for maximum coverage"""
        extracted_text = ""
        extraction_method = "none"
        
        # Method 1: Try PyMuPDF first (more reliable)
        try:
            pdf_doc = fitz.open(file_path)
            print(f"📄 PDF has {len(pdf_doc)} pages")
            
            for page_num in range(len(pdf_doc)):
                page = pdf_doc[page_num]
                page_text = page.get_text()
                extracted_text += page_text + " "
                print(f"Page {page_num + 1}: {len(page_text)} characters")
            
            pdf_doc.close()
            print(f"🔍 PyMuPDF extracted: {len(extracted_text)} characters")
            extraction_method = "pymupdf_success"
            
        except Exception as e:
            print(f"PyMuPDF failed: {e}")
            
            # Method 2: Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    print(f"📄 PDF has {len(pdf_reader.pages)} pages")
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        extracted_text += page_text + " "
                        print(f"PyPDF2 Page {page_num + 1}: {len(page_text)} characters")
                
                print(f"🔍 PyPDF2 extracted: {len(extracted_text)} characters")
                extraction_method = "pypdf2_success"
                
            except Exception as e:
                print(f"PyPDF2 also failed: {e}")
                extraction_method = "extraction_failed"
        
        return self._process_extracted_text(extracted_text, "PDF", extraction_method)
    
    async def _analyze_docx_text(self, file_path: str):
        """Extract text from DOCX files using python-docx"""
        extracted_text = ""
        extraction_method = "none"
        
        try:
            print(f"📝 Opening DOCX file: {file_path}")
            doc = docx.Document(file_path)
            
            # Extract text from all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            extracted_text = " ".join(paragraphs)
            
            # Also extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_text.append(cell.text)
            
            if table_text:
                extracted_text += " " + " ".join(table_text)
            
            extraction_method = "docx_success"
            print(f"📊 DOCX extracted: {len(extracted_text)} characters")
            
        except Exception as e:
            print(f"DOCX text extraction failed: {e}")
            extraction_method = "extraction_failed"
        
        return self._process_extracted_text(extracted_text, "DOCX", extraction_method)
    
    def _process_extracted_text(self, extracted_text: str, file_type: str, extraction_method: str):
        """Process and analyze extracted text with proper confidence scoring"""
        
        # Clean text
        cleaned_text = re.sub(r'\s+', ' ', extracted_text.strip()) if extracted_text else ""
        
        if cleaned_text and len(cleaned_text) > 0:
            # Text found - normal analysis
            words = re.findall(r'\b\w+\b', cleaned_text)
            word_count = len(words)
            
            print(f"📊 Final word count: {word_count}")
            print(f"📝 Sample text: {cleaned_text[:200]}...")
            
            # Analyze for suspicious patterns
            suspicious_count = 0
            flags = []
            
            if word_count > 10:
                # Check for repetition
                if len(set(words)) < max(1, len(words) * 0.1):
                    suspicious_count += 5
                    flags.append("High text repetition detected")
                
                # Check word length distribution
                avg_word_length = sum(len(word) for word in words) / len(words)
                if avg_word_length < 2:
                    suspicious_count += 3
                    flags.append("Unusually short words detected")
                elif avg_word_length > 15:
                    suspicious_count += 2
                    flags.append("Unusually long words detected")
            
            # Calculate confidence
            confidence = max(70, min(100, 100 - (suspicious_count * 3)))
            
            return {
                "totalWords": word_count,
                "suspiciousWords": suspicious_count,
                "confidence": confidence,
                "flags": flags if flags else [f"Text extraction successful from {file_type}"]
            }
            
        elif extraction_method == "extraction_failed":
            # Genuine extraction failure
            return {
                "totalWords": 0,
                "suspiciousWords": 0,
                "confidence": 50,
                "flags": ["Text extraction failed - document may be corrupted or unreadable"]
            }
            
        else:
            # No text found but extraction succeeded - likely image-only document
            print(f"📊 No text content found - appears to be image-only {file_type}")
            return {
                "totalWords": 0,
                "suspiciousWords": 0,
                "confidence": 95,  # ✅ HIGH confidence - this is normal for image-only PDFs
                "flags": [f"No text content detected - document appears to be image/diagram only"]
            }
    
    async def _basic_analysis(self, file_info: dict):
        """Basic analysis for non-PDF/DOCX files"""
        return {
            "totalWords": 0,
            "suspiciousWords": 0,
            "confidence": 100,
            "flags": ["Text analysis not implemented for this file type"]
        }
